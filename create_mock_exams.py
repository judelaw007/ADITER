#!/usr/bin/env python3
"""
Generate 10 Mock Exams for ADIT Module 3.03 - Transfer Pricing
Following the exam format: Part A (Case Studies), Part B (Scenario), Part C (Short Questions)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os

def add_heading(doc, text, level=1):
    """Add a heading to the document"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.bold = True
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_paragraph(doc, text, bold=False, italic=False, size=11):
    """Add a paragraph to the document"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
    return p

def create_exam_header(doc, exam_number):
    """Create the exam header"""
    doc.add_heading(f'MOCK EXAMINATION', level=0)
    doc.add_heading(f'ADIT Module 3.03 – Transfer Pricing Option', level=2)
    doc.add_heading(f'Mock Exam {exam_number}', level=2)

    # Exam Details
    details = [
        'Time Allowed: 3 hours 15 minutes',
        'Total Marks: 100 marks',
        f'Date: {datetime.now().strftime("%B %Y")}',
        '',
        'INSTRUCTIONS TO CANDIDATES:',
        '• Answer FIVE questions in total',
        '• PART A: Answer BOTH questions (2 × 25 marks = 50 marks)',
        '• PART B: Answer ONE question from the two offered (20 marks)',
        '• PART C: Answer TWO questions from the five offered (2 × 15 marks = 30 marks)',
        '',
        'Write your answers clearly and legibly. Marks will be awarded for:',
        '• Correct identification and analysis of issues',
        '• Application of OECD Transfer Pricing Guidelines',
        '• Supporting calculations and references',
        '• Clear presentation and structure',
    ]

    for detail in details:
        if detail == '':
            doc.add_paragraph()
        else:
            add_paragraph(doc, detail)

    doc.add_page_break()

def create_part_a(doc, exam_num):
    """Create Part A - Compulsory Case Studies"""
    add_heading(doc, 'PART A: COMPULSORY CASE STUDY QUESTIONS', level=1)
    add_paragraph(doc, 'Answer BOTH questions. Each question is worth 25 marks.', italic=True)
    doc.add_paragraph()

    if exam_num == 1:
        create_exam1_part_a(doc)
    elif exam_num == 2:
        create_exam2_part_a(doc)
    elif exam_num == 3:
        create_exam3_part_a(doc)
    elif exam_num == 4:
        create_exam4_part_a(doc)
    elif exam_num == 5:
        create_exam5_part_a(doc)
    elif exam_num == 6:
        create_exam6_part_a(doc)
    elif exam_num == 7:
        create_exam7_part_a(doc)
    elif exam_num == 8:
        create_exam8_part_a(doc)
    elif exam_num == 9:
        create_exam9_part_a(doc)
    elif exam_num == 10:
        create_exam10_part_a(doc)

def create_exam1_part_a(doc):
    """Mock Exam 1 - Global Tech Solutions (Software Development and IP Licensing)"""
    add_heading(doc, 'QUESTION 1 (25 marks)', level=2)

    scenario = """
SCENARIO - Global Tech Solutions Group

Global Tech Solutions is an international group specializing in software development and cloud computing solutions.
The group was established in 2020 and has the following structure:

1. GTS Parent (United States) - Parent company and IP owner
2. GTS Europe (Ireland) - European service hub
3. GTS APAC (Singapore) - Asia-Pacific operations
4. GTS India (India) - Development center
5. GTS Brazil (Brazil) - Software development and support

KEY FACTS:

GTS Parent:
- Holds all intellectual property rights (patents, software, databases, brand)
- Provides strategic planning and group coordination
- Employs 20 executives and IT specialists
- Assets: USD 50 million (primarily IP)

GTS Europe (Ireland):
- Operates as the European service hub
- Provides consulting services, implementation services, and customer support
- Manages sales to European clients
- Employs 150 employees (consultants, developers, support staff)
- Annual revenue: EUR 60 million
- Pays 12.5% corporate tax rate

GTS APAC (Singapore):
- Regional headquarters for Asia-Pacific
- Limited development activities
- Primarily coordination, project management, and customer relationship management
- Employs 30 employees
- Annual revenue: SGD 40 million
- Tax rate: 17%

GTS India (Development Center):
- Provides software development and technical support services
- Employs 300 developers and technical staff
- Receives assignments from GTS Parent and GTS Europe
- Provides pre-developed code modules, custom solutions, and IT support
- Cost of operations: INR 250 million per annum
- Tax rate: 30%

GTS Brazil (Support Center):
- Recent establishment (2023) for redundancy and risk management
- Provides software maintenance and customer support
- Employs 50 support staff
- Annual costs: BRL 15 million
- Tax rate: 34%

TRANSACTIONS AND FINANCIAL DATA:

A. Intragroup Transactions:
   1. GTS Parent licenses software and IP to GTS Europe and GTS APAC for:
      - GTS Europe: Annual royalty USD 8 million
      - GTS APAC: Annual royalty USD 2 million

   2. GTS India provides development services to:
      - GTS Parent: USD 5 million per annum (cost-reimbursement basis)
      - GTS Europe: USD 15 million per annum (cost-reimbursement basis)

   3. GTS Europe charges:
      - Implementation and consulting services to clients: average charge rate EUR 200/hour
      - Service cost recharge to GTS India: EUR 2 million per annum for project management
      - Internal allocation of corporate costs

   4. GTS Brazil provides support services:
      - To GTS Europe: BRL 3 million per annum
      - To GTS Parent: BRL 2 million per annum

B. Group Financials (Audited):
   - GTS Parent: Net profit USD 25 million on revenue USD 40 million (62.5% margin)
   - GTS Europe: Net profit EUR 3 million on revenue EUR 60 million (5% margin)
   - GTS APAC: Net profit SGD 1 million on revenue SGD 40 million (2.5% margin)
   - GTS India: Net loss INR 20 million on revenue INR 250 million (cost reimbursement)
   - GTS Brazil: Loss BRL 2 million on costs BRL 15 million

ADDITIONAL INFORMATION:
- Market data shows comparable software companies charge 15-25% royalty rates for standard IP licenses
- Development centers in India typically charge cost-plus 10-15% for routine development
- European IT consultants typically achieve net margins of 10-15%
- No transfer pricing documentation or APAs are in place
- The group has not been subject to any transfer pricing audits
- All entities operate at arm's length as independent entities would

REQUIRED:

(a) Identify and describe all controlled transactions within the Global Tech Solutions group.
    What additional information would you require to properly characterize these transactions?
    (8 marks)

(b) Perform a detailed functional analysis for each entity, identifying:
    - Functions performed
    - Assets employed
    - Risks assumed
    - Business justification for each operation
    (10 marks)

(c) For the IP licensing arrangements, evaluate whether the current royalty rates are
    arm's length under the OECD Transfer Pricing Guidelines. What transfer pricing method
    would be most appropriate and why?
    (7 marks)
"""
    add_paragraph(doc, scenario)

    add_heading(doc, 'QUESTION 2 (25 marks)', level=2)

    question2 = """
With reference to the Global Tech Solutions scenario above, answer the following:

(a) Characterize each entity (GTS Europe, GTS APAC, GTS India, and GTS Brazil) according to
    OECD transfer pricing classifications. Do you believe the current profit allocation is
    consistent with the arm's length principle? Explain your analysis.
    (10 marks)

(b) The operating margins achieved by each entity vary significantly. Discuss whether these
    differences can be justified by differences in functions, assets, and risks. Are there
    any transfer pricing concerns from a tax authority perspective?
    (8 marks)

(c) Identify specific transfer pricing risks in this structure. What advice would you
    provide to the group regarding documentation and compliance?
    (7 marks)
"""
    add_paragraph(doc, question2)
    doc.add_page_break()

def create_exam2_part_a(doc):
    """Mock Exam 2 - Luxe Fashion Group (Manufacturing and Distribution)"""
    add_heading(doc, 'QUESTION 1 (25 marks)', level=2)

    scenario = """
SCENARIO - Luxe Fashion Group

Luxe Fashion Group is a multinational luxury apparel company with operations across multiple
jurisdictions. The group manufactures and sells high-end fashion products globally.

GROUP STRUCTURE:

1. Luxe Holdings (Luxembourg) - Parent company, brand owner, IP holder
2. Luxe Manufacturing (Vietnam) - Contract manufacturer
3. Luxe Distribution Europe (Germany) - European distribution hub
4. Luxe Distribution Asia (Hong Kong) - Asian distribution center
5. Luxe Retail (UK) - Retail operations and UK subsidiary

COMPANY PROFILES:

Luxe Holdings (Luxembourg):
- Ultimate parent company
- Owns the "Luxe" brand and all associated trademarks
- Owns product designs and patterns
- Provides marketing strategies and brand management
- Minimal direct involvement in production or distribution
- Employees: 8 (executives and brand managers)
- Assets: Brand value (estimated USD 200 million), designs
- Tax rate: 0.29% (Luxembourg tax incentive)

Luxe Manufacturing (Vietnam):
- Contract manufacturer of luxury apparel
- Receives designs and specifications from Luxe Holdings
- Sources raw materials and components
- Manufactures products to specified standards
- No intellectual property ownership
- Limited marketing involvement
- Employees: 500
- Annual production capacity: 1 million units
- Current utilization: 80%
- Cost of goods manufactured (COGS): USD 50 per unit
- Operating costs: USD 15 million per annum
- Products sold to distribution centers at cost-plus 15% pricing

Luxe Distribution Europe (Germany):
- Receives finished goods from manufacturing
- Warehouses and distributes to retail partners across Europe
- Conducts marketing campaigns (local adaptation)
- Handles customer relationships and warranty claims
- Employees: 150
- Assets: Warehouse facility (value USD 10 million), brand awareness developed locally
- Annual purchases from manufacturer: 400,000 units
- Purchasing cost: USD 57.50 per unit (cost-plus 15%)
- Selling price to retailers: USD 150 per unit
- Operating expenses: EUR 8 million per annum

Luxe Distribution Asia (Hong Kong):
- Similar role to Europe distribution entity
- Distributes to retailers across Asia-Pacific
- Less established market presence compared to Europe
- Employees: 50
- Annual purchases: 200,000 units
- Purchasing cost: USD 57.50 per unit
- Average selling price: USD 140 per unit
- Operating expenses: HKD 20 million per annum

Luxe Retail (UK):
- Operates flagship retail stores in major UK cities
- Direct consumer sales
- Brand experience and customer relationship management
- Employees: 120
- Annual unit sales: 100,000 units
- Purchasing cost from Distribution Europe: USD 80 per unit
- Average retail price: USD 250 per unit
- Operating expenses: GBP 5 million per annum
- Store lease and utilities: GBP 3 million per annum

ADDITIONAL CONSIDERATIONS:
- Manufacturing labor costs in Vietnam are significantly lower than in developed countries
- The "Luxe" brand has strong recognition in European and Asian markets
- Retail customer base is primarily high-net-worth individuals
- Product quality and brand positioning are critical to pricing power
- No comparables for exact brand or product type available
- Recent market analysis shows luxury apparel companies achieve 40-60% retail margins

REQUIRED:

(a) Delineate all controlled transactions in the Luxe Fashion Group structure. For each
    transaction, identify the nature, parties, and economic context.
    (8 marks)

(b) Evaluate the current transfer pricing arrangements (cost-plus 15% for manufacturing).
    Is this methodology appropriate? What alternative methodologies might be considered
    and why?
    (9 marks)

(c) What comparability data would you require to support the transfer pricing positions?
    Discuss the challenges in finding relevant comparables for this group.
    (8 marks)
"""
    add_paragraph(doc, scenario)

    add_heading(doc, 'QUESTION 2 (25 marks)', level=2)

    question2 = """
Continuing with the Luxe Fashion Group scenario:

(a) Characterize each entity in the value chain. Identify which entity should logically
    earn the residual profit and why, based on OECD transfer pricing principles.
    (10 marks)

(b) The retail operations achieve significantly higher profit margins than distribution
    centers. Analyze whether this is consistent with the arm's length principle. Are there
    transfer pricing concerns?
    (8 marks)

(c) Recommend an alternative transfer pricing structure for the group that you believe would
    be more defensible. Justify your recommendation with reference to OECD Guidelines.
    (7 marks)
"""
    add_paragraph(doc, question2)
    doc.add_page_break()

def create_exam3_part_a(doc):
    """Mock Exam 3 - Global Finance Holdings (Intragroup Financing)"""
    add_heading(doc, 'QUESTION 1 (25 marks)', level=2)

    scenario = """
SCENARIO - Global Finance Holdings

Global Finance Holdings is a multinational holding company with significant financing
operations among group companies. The group structure involves cross-border loans and
intercompany guarantees.

GROUP STRUCTURE:

1. Global Finance Holding (Netherlands) - Ultimate parent, credit rating A-
2. Operating Company A (France) - Operating subsidiary
3. Operating Company B (Brazil) - Operating subsidiary
4. Finance Company (Luxembourg) - Finance subsidiary
5. Start-Up Company (Poland) - Recently established entity

FINANCIAL DETAILS:

Global Finance Holding (Netherlands):
- Investment holding company
- Lends to group companies
- Credit rating: A- (investment grade)
- Annual borrowing rate available: 2.5%
- Cash resources: EUR 500 million
- Outstanding loans to subsidiaries: EUR 300 million

Finance Company (Luxembourg):
- Central financing vehicle
- Raises funds on capital markets
- On-lends to group operating companies
- Cost of debt: 2% (treasury bills rate)
- Borrowing from parent: 2.5%
- On-lending rates to Operating Company A: 4%
- On-lending rates to Operating Company B: 5.5%

Operating Company A (France):
- Manufacturing operations
- Loan from Finance Company: EUR 100 million at 4%
- Loan balance: EUR 100 million
- EBIT: EUR 20 million
- Tax rate: 25%
- Debt capacity: EUR 150 million
- Credit rating: BB+ (sub-investment grade)

Operating Company B (Brazil):
- Mining operations
- Loan from Finance Company: EUR 80 million at 5.5%
- Loan balance: EUR 80 million
- EBIT: EUR 25 million
- Tax rate: 34%
- Debt capacity: EUR 100 million
- Credit rating: BB (sub-investment grade)
- Volatile earnings (commodity dependent)

Start-Up Company (Poland):
- E-commerce venture established 2023
- Loan from Global Finance Holding: EUR 50 million at 3%
- Loan balance: EUR 50 million
- Annual losses: EUR 5 million
- Tax rate: 19%
- No independent debt capacity
- Credit rating: Not rated (no history)

BACKGROUND INFORMATION:
- A-rated companies can borrow at 2.5% in current market conditions
- BB+ rated companies typically pay 4-5% interest
- BB rated companies typically pay 5-6% interest
- Start-up companies with no history cannot access independent financing
- Market analysis shows an average spread of 200-300 basis points over parent rate for
  sub-investment grade subsidiaries
- No documentation exists for these financing arrangements
- No APAs are in place

REQUIRED:

(a) For each intragroup loan, evaluate whether the interest rate charged is arm's length:
    (i)   Global Finance Holding to Finance Company (2.5%)
    (ii)  Finance Company to Operating Company A (4%)
    (iii) Finance Company to Operating Company B (5.5%)
    (iv)  Global Finance Holding to Start-Up Company (3%)

    Explain your analysis for each rate with reference to the OECD Transfer Pricing
    Guidelines Section X (Financial Transactions).
    (12 marks)

(b) Discuss any transfer pricing risks or concerns regarding:
    - The implicit guarantees provided by the parent company
    - The relationship between the interest rate and the borrowing subsidiary's creditworthiness
    - Tax optimization potential of the financing structure
    (8 marks)

(c) What additional information or documentation would you recommend the group prepare
    to support these financing arrangements?
    (5 marks)
"""
    add_paragraph(doc, scenario)

    add_heading(doc, 'QUESTION 2 (25 marks)', level=2)

    question2 = """
Continuing with the Global Finance Holdings scenario:

(a) The Start-Up Company in Poland is unable to access independent market financing.
    Analyze the transfer pricing implications and discuss how the group should structure
    the financing to be defensible.
    (10 marks)

(b) Brazil is considering increasing its leverage to take advantage of the lower tax rate
    (34%) compared to France (25%). The Finance Company has proposed increasing the loan
    balance to EUR 150 million at 6% interest. Evaluate the transfer pricing and tax
    issues associated with this proposal.
    (8 marks)

(c) Recommend a sustainable financing structure that balances commercial rationale with
    transfer pricing compliance across all group entities.
    (7 marks)
"""
    add_paragraph(doc, question2)
    doc.add_page_break()

def create_exam4_part_a(doc):
    """Mock Exam 4 - Digital Services Group (Technology and PE)"""
    add_heading(doc, 'QUESTION 1 (25 marks)', level=2)

    scenario = """
SCENARIO - Digital Services Group

Digital Services Group provides cloud computing, data analytics, and IT consulting services
globally. The group operates through a complex structure with operations in multiple
jurisdictions.

GROUP STRUCTURE:

1. Digital Holdings (Switzerland) - Parent company
2. Digital Development (Ireland) - Software development and product development
3. Digital Services (US) - Service delivery and implementation center
4. Digital EMEA (UK) - Customer service hub for Europe, Middle East, Africa
5. Digital Operations (India) - IT operations and service delivery

KEY OPERATIONS:

Digital Holdings (Switzerland):
- Ultimate parent company
- Strategic planning and group direction
- IP ownership and licensing
- Treasury function
- Employees: 15 executives
- Tax rate: 13.2%
- Owns key technologies and cloud infrastructure patents

Digital Development (Ireland):
- Develops and maintains cloud platform software
- Creates data analytics tools
- Conducts R&D activities
- Employs 200 developers and data scientists
- Annual development costs: EUR 30 million
- Licenses platform to group companies
- Annual platform license fee to services companies: EUR 10 million
- Tax rate: 12.5%

Digital Services (US):
- Implements cloud solutions for US corporate clients
- Provides consulting and system integration services
- Direct sales to end customers in US market
- Employees: 300 consultants and implementation staff
- Annual service revenue: USD 200 million
- On-shores work from external providers
- Tax rate: 21%

Digital EMEA (UK):
- Service delivery and customer support for EMEA region
- Handles implementation and managed services
- Employees: 150
- Annual service revenue: GBP 60 million
- Pays platform license fee: GBP 6 million
- Tax rate: 25%

Digital Operations (India):
- Provides offshore IT operations and technical support
- Delivers platform implementation and maintenance
- Employs 400 IT specialists
- Annual revenue: INR 500 million
- Cost-reimbursement basis from other group entities
- Annual costs: INR 450 million
- Tax rate: 30%

TRANSACTIONS:
- Digital Development charges 10% of service revenues as platform license fees
- Digital Operations provides services on cost-plus 10% basis
- All entities use services from Digital Operations

REQUIRED:

(a) Identify and describe all controlled transactions. For the platform licensing
    arrangement, evaluate the appropriateness of the 10% royalty rate.
    (10 marks)

(b) Perform a detailed functional analysis for Digital Services (US) and Digital EMEA (UK),
    identifying key differences and similarities in their operations and risk profiles.
    (8 marks)

(c) What transfer pricing method would you recommend for the platform licensing arrangement?
    What comparables or benchmarking would you require?
    (7 marks)
"""
    add_paragraph(doc, scenario)

    add_heading(doc, 'QUESTION 2 (25 marks)', level=2)

    question2 = """
Continuing with Digital Services Group:

(a) The UK tax authority is concerned about the structure and believes Digital EMEA should
    be paying higher royalty rates. Perform a comparable analysis and recommend an
    appropriate royalty rate range with supporting justification.
    (12 marks)

(b) Discuss potential Permanent Establishment issues if:
    (i)   Digital Services maintains a project team on-site at a major UK customer location
          for 24 months
    (ii)  Digital EMEA coordinates implementation services conducted by Digital Operations
          personnel in client offices across EMEA
    (9 marks)

(c) Identify transfer pricing documentation that the group should prepare to support
    its positions in the US, UK, and India.
    (4 marks)
"""
    add_paragraph(doc, question2)
    doc.add_page_break()

def create_exam5_part_a(doc):
    """Mock Exam 5 - Pharma Corp (IP and R&D)"""
    add_heading(doc, 'QUESTION 1 (25 marks)', level=2)

    scenario = """
SCENARIO - Pharma Corp

Pharma Corp is a multinational pharmaceutical company with significant research and
development operations and IP licensing arrangements.

GROUP STRUCTURE:

1. Pharma International (Switzerland) - Parent and IP holder
2. Pharma Research (Canada) - R&D center
3. Pharma Manufacturing (Puerto Rico) - Manufacturing
4. Pharma Distribution (UK) - European distribution
5. Pharma Markets (Japan) - Japanese affiliate

COMPANY DETAILS:

Pharma International (Switzerland):
- Ultimate parent company
- Owns all patents and pharmaceutical IP
- Coordinates group strategy
- Conducts regulatory liaison
- Employees: 20
- Tax rate: 12%
- IP portfolio value: USD 1.2 billion

Pharma Research (Canada):
- Conducts drug development and clinical trials
- Performs pre-clinical and clinical research
- Employees: 150 scientists and technicians
- Annual R&D costs: CAD 100 million
- No direct drug sales
- Conducts research on cost-reimbursement basis
- Tax rate: 26%
- Awarded research tax credits offsetting costs

Pharma Manufacturing (Puerto Rico):
- Manufactures approved pharmaceutical products
- Sources raw materials
- Quality control and packaging
- Employees: 200
- Manufacturing volume: 50 million units per annum
- Cost of manufacturing: USD 15 per unit (including COGS and direct labor)
- Operating expenses: USD 10 million per annum
- Sells to distribution entities at cost-plus 20%
- Tax rate: 4% (Act 60 tax incentive - Export Services)

Pharma Distribution (UK):
- Distributes pharmaceutical products across Europe
- Handles regulatory compliance and marketing
- Manages customer relationships
- Employees: 100
- Annual unit purchases: 30 million units
- Purchase price: USD 18 per unit (cost-plus 20%)
- Average selling price: USD 60 per unit
- Marketing and distribution costs: GBP 15 million per annum
- Tax rate: 19%

Pharma Markets (Japan):
- Distributes products in Japanese market
- Conducts extensive market development
- Employees: 80
- Annual unit purchases: 10 million units
- Purchase price: USD 18 per unit
- Average selling price: JPY 3,000 per unit (USD equivalent: USD 20)
- Operating expenses: JPY 2 billion per annum
- Tax rate: 30%

IP LICENSING:
- Pharma International licenses all IP to manufacturing and distribution entities
- Current licensing arrangement: royalty based on units manufactured (USD 2 per unit)
- Total annual royalty revenue: USD 120 million

ADDITIONAL INFORMATION:
- Pharma International claims only 15% of R&D costs (rest borne by Pharma Research)
- BEPS Actions 8-10 significantly tightened DEMPE analysis for pharmaceutical companies
- Market data shows pharmaceutical companies typically charge 5-8% royalty rates on sales
- Pharma Markets operates at a loss in early years of market development
- Patent protection in various jurisdictions expires at different times

REQUIRED:

(a) Analyze the DEMPE functions (Development, Enhancement, Maintenance, Protection,
    Exploitation) across the group:
    (i)   Who performs each function?
    (ii)  Who bears the costs?
    (iii) Who assumes risks?
    (iv)  Are the functions appropriately allocated among group entities?
    (10 marks)

(b) Evaluate whether the current IP licensing arrangement (USD 2 per unit) is arm's length.
    Apply OECD Transfer Pricing Guidelines provisions on IP to your analysis.
    (9 marks)

(c) What transfer pricing issues do you identify with this structure from a BEPS perspective?
    (6 marks)
"""
    add_paragraph(doc, scenario)

    add_heading(doc, 'QUESTION 2 (25 marks)', level=2)

    question2 = """
Continuing with Pharma Corp:

(a) The Pharma Markets (Japan) entity is operating at a loss. Analyze whether this is
    consistent with the arm's length principle. Should the transfer pricing be adjusted?
    What factors should be considered?
    (10 marks)

(b) Discuss the transfer pricing implications of Pharma International recognizing only 15%
    of R&D costs. Is this allocation defensible under OECD Guidelines and BEPS Actions
    8-10?
    (9 marks)

(c) Recommend a revised IP licensing arrangement (including the royalty rate and structure)
    that would be more robust to challenge. Provide supporting analysis.
    (6 marks)
"""
    add_paragraph(doc, question2)
    doc.add_page_break()

def create_exam6_part_a(doc):
    """Mock Exam 6 - E-Commerce Plus (Digital Platform and PE)"""
    add_heading(doc, 'QUESTION 1 (25 marks)', level=2)

    scenario = """
SCENARIO - E-Commerce Plus

E-Commerce Plus is an online marketplace platform operating across multiple countries.
The group includes entity structures that raise potential PE and transfer pricing issues.

GROUP STRUCTURE:

1. E-Commerce International (Netherlands) - Parent platform owner
2. E-Commerce Operations (Germany) - EU operations and customer service
3. E-Commerce Fulfillment (Poland) - Warehouse and logistics
4. E-Commerce Enabler (China) - Technology platform development
5. E-Commerce Markets (Australia) - Local market operations

KEY OPERATIONS:

E-Commerce International (Netherlands):
- Owns the e-commerce platform
- Sets group strategy and policies
- Owns brand and technology IP
- Provides customer acquisition strategy
- Employees: 30
- Tax rate: 15%

E-Commerce Operations (Germany):
- Manages EU customer relationships
- Handles customer service (multiple languages)
- Conducts targeted marketing campaigns (local)
- Marketing database and customer analytics
- Employees: 200
- Annual customer service costs: EUR 25 million
- Annual marketing costs: EUR 15 million
- Annual revenue recognition: EUR 5 million (service fee from platform)
- Tax rate: 30%

E-Commerce Fulfillment (Poland):
- Operates warehouse facilities
- Receives, stores, and ships inventory
- Quality control on fulfillment
- Employees: 150
- Annual facility operating costs: PLN 50 million
- Storage and handling fees from sellers: PLN 35 million
- Margin on fulfillment operations: PLN 8 million
- Tax rate: 19%

E-Commerce Enabler (China):
- Develops and maintains the platform technology
- Provides infrastructure and hosting
- Conducts platform improvements and feature development
- Employees: 100
- Annual development and technology costs: CNY 150 million
- Charges platform maintenance fee: CNY 50 million per annum
- Tax rate: 25%

E-Commerce Markets (Australia):
- Established 18 months ago
- Markets the platform in Australia
- Handles disputes and local compliance
- Has not yet generated significant sales
- Employees: 15
- Annual costs: AUD 3 million
- Annual revenue: AUD 0.1 million
- Tax rate: 30%

OPERATING MODEL:
- Individual sellers list products on platform
- E-Commerce International owns the platform technology (licensed from Enabler)
- E-Commerce Operations manages EU customer acquisition
- E-Commerce Fulfillment provides optional fulfillment services
- E-Commerce Markets developing Australian market presence
- Fee structure: Commission on sales (30%), platform maintenance fee, fulfillment fees

REQUIRED:

(a) Delineate the controlled transactions in the E-Commerce Plus structure. Identify the
    economic substance and transfer pricing treatment for each transaction.
    (8 marks)

(b) Perform a functional analysis for E-Commerce Operations and E-Commerce Markets.
    Discuss whether their different profit levels can be explained by differences in
    functions and risks.
    (9 marks)

(c) Identify and discuss potential Permanent Establishment issues arising from the current
    structure. Consider whether any entity might constitute a PE in any jurisdiction.
    (8 marks)
"""
    add_paragraph(doc, scenario)

    add_heading(doc, 'QUESTION 2 (25 marks)', level=2)

    question2 = """
Continuing with E-Commerce Plus:

(a) The Australian Tax Office has challenged whether E-Commerce Markets constitutes a PE
    of E-Commerce International. Analyze this issue thoroughly, considering:
    (i)   Whether E-Commerce Markets has a "fixed place of business"
    (ii)  The significance of its current operations
    (iii) How profits should be attributed if a PE exists
    (10 marks)

(b) Recommend a sustainable transfer pricing structure for:
    (i)   The platform technology licensing arrangement
    (ii)  The customer service and marketing functions in Germany
    (iii) The fulfillment operations in Poland
    Justify your recommendations with reference to OECD Guidelines.
    (10 marks)

(c) What transfer pricing documentation should the group prepare to support its positions
    under the BEPS Action 13 three-tiered approach?
    (5 marks)
"""
    add_paragraph(doc, question2)
    doc.add_page_break()

def create_exam7_part_a(doc):
    """Mock Exam 7 - Construction Global (Services and PE)"""
    add_heading(doc, 'QUESTION 1 (25 marks)', level=2)

    scenario = """
SCENARIO - Construction Global

Construction Global is a multinational construction and engineering services company
operating on major infrastructure projects across continents.

GROUP STRUCTURE:

1. Construction Global Holding (Luxembourg) - Parent company
2. Construction Europe (Belgium) - European operations
3. Construction Americas (USA) - Americas operations
4. Construction Engineering (India) - Design and engineering
5. Construction Resources (UAE) - Project management and coordination

COMPANY PROFILES:

Construction Global Holding (Luxembourg):
- Parent company providing group coordination
- Bid on large international projects
- Provide management and oversight
- Employees: 25 management and business development
- Tax rate: 17%

Construction Europe (Belgium):
- Executes construction projects in Europe
- Brings in engineering from India
- Project team on-site at customer locations
- Employees: 200 project managers, site supervisors, some labor
- Annual revenue from projects: EUR 150 million
- Annual direct labor and costs: EUR 120 million
- Annual net profit: EUR 15 million
- Tax rate: 29.58%
- Typically on-site for 18-36 months per project

Construction Americas (USA):
- Executes construction projects in North and South America
- Brings in engineering from India
- Project team on-site at customer locations
- Employees: 300 project managers, engineers, some labor
- Annual revenue from projects: USD 200 million
- Annual costs: USD 170 million
- Annual net profit: USD 20 million
- Tax rate: 21%
- Current major project: 28-month infrastructure project in Mexico

Construction Engineering (India):
- Provides engineering design and technical services
- Provides key technical personnel for project teams
- Performs site inspections and technical oversight
- Employees: 150 engineers and technical specialists
- Annual revenue: INR 350 million (for engineering and technical services)
- Annual costs: INR 310 million
- Profit margin: INR 40 million (11%)
- Tax rate: 30%

Construction Resources (UAE):
- Project management office for Middle East and Africa
- Central coordination and planning
- Equipment storage and logistics
- Employees: 80
- Annual coordination and management fees: AED 50 million
- Annual costs: AED 30 million
- Profit: AED 20 million
- Tax rate: 0% (no corporate tax in UAE)

PROJECT STRUCTURE:
Example: 28-Month Construction Project in Mexico

Total project value: USD 200 million
- Construction Europe wins bid: USD 80 million portion
- Construction Americas wins bid: USD 120 million portion
- Construction Engineering provides technical services: INR 1.2 billion (USD 14 million)
- Construction Resources coordinates: AED 30 million (USD 8 million)

Key personnel on-site in Mexico for 28 months:
- From Construction Europe: 50 employees
- From Construction Americas: 80 employees
- From Construction Engineering: 20 key technical staff members
- From Construction Resources: 5 project coordinators

REQUIRED:

(a) Analyze whether Construction Engineering's activities in Mexico constitute a Permanent
    Establishment under Article 5 of the OECD Model Tax Convention. Specifically consider:
    (i)   Whether there is a "fixed place of business"
    (ii)  The duration and nature of the presence
    (iii) The dependent agent PE rules
    (10 marks)

(b) If a PE is deemed to exist, discuss how profits should be attributed to the PE using
    the Authorized OECD Approach. What functions and assets would you attribute to the PE?
    (9 marks)

(c) Evaluate the current transfer pricing arrangement for engineering services (INR 1.2
    billion). Is this arm's length? What factors should influence the pricing?
    (6 marks)
"""
    add_paragraph(doc, scenario)

    add_heading(doc, 'QUESTION 2 (25 marks)', level=2)

    question2 = """
Continuing with Construction Global:

(a) Discuss transfer pricing issues for the services provided by Construction Resources
    (UAE). Is the profit allocation appropriate given the UAE has no corporate tax? What
    risks or concerns should the group consider?
    (10 marks)

(b) The Mexico tax authority has challenged the construction group's transfer pricing,
    claiming that Construction Engineering should bear greater costs given the nature of
    its activities. Respond to this challenge with a complete analysis and recommendation.
    (10 marks)

(c) Recommend changes to the group's structure or transfer pricing arrangements to
    minimize the risk of PE challenges in future projects while maintaining operational
    efficiency.
    (5 marks)
"""
    add_paragraph(doc, question2)
    doc.add_page_break()

def create_exam8_part_a(doc):
    """Mock Exam 8 - Artisan Products (Business Restructuring)"""
    add_heading(doc, 'QUESTION 1 (25 marks)', level=2)

    scenario = """
SCENARIO - Artisan Products Restructuring

Artisan Products is an international luxury goods manufacturer that is undertaking a
significant business restructuring to centralize manufacturing and optimize the supply chain.

CURRENT STRUCTURE (Before Restructuring):

1. Artisan Group (Switzerland) - Parent company
2. Artisan Manufacturing EU (Italy) - Manufacturing for EU
3. Artisan Manufacturing Asia (Vietnam) - Manufacturing for Asia
4. Artisan Distribution USA (USA) - USA distribution and sales
5. Artisan Distribution EU (France) - EU distribution and sales

CURRENT OPERATIONS (Pre-Restructuring):

Artisan Manufacturing EU (Italy):
- Manufactures products for EU market
- Product design and development
- Quality control
- Employees: 250
- Annual production: 300,000 units
- Manufacturing cost: EUR 60 per unit
- EBIT: EUR 15 million
- Tax rate: 24%
- Owns local manufacturing relationships and supply chain expertise

Artisan Manufacturing Asia (Vietnam):
- Manufactures products for Asian markets
- Lower-cost production facility
- Employees: 400
- Annual production: 500,000 units
- Manufacturing cost: USD 45 per unit
- EBIT: USD 25 million
- Tax rate: 20%
- Local government incentives and relationships

Artisan Distribution USA (USA):
- Manages US sales and distribution
- Brand marketing in US
- High-end retail relationships
- Employees: 150
- Annual sales: 200,000 units at USD 250 per unit
- Annual costs: USD 15 million
- Net profit: USD 35 million
- Tax rate: 21%

Artisan Distribution EU (France):
- Manages EU distribution
- Marketing and brand building
- Retailer relationships
- Employees: 100
- Annual sales: 200,000 units at EUR 200 per unit
- Annual costs: EUR 12 million
- Net profit: EUR 28 million
- Tax rate: 28%

PROPOSED RESTRUCTURING:

The group intends to:

1. Consolidate all manufacturing into Vietnam facility
   - Close Italy manufacturing facility by end of Year 1
   - Transfer manufacturing knowledge and supply chains to Vietnam
   - Eliminate EU-based production costs

2. Retain Design and IP with Artisan Group (Switzerland)
   - Continue product design and development
   - IP ownership remains with parent
   - Licensing of designs to Vietnam manufacturer

3. Establish "Principal" Distributor Model
   - Both USA and EU distribution entities assume principal distributor role
   - Greater responsibility for customer acquisition, brand building, product returns
   - Larger inventory holdings and working capital

4. Manufacturing Transition
   - Vietnam facility to assume all manufacturing
   - Increased production volume to 800,000 units
   - Scale economies expected to reduce costs to USD 40 per unit

REQUIRED:

(a) For the Italy manufacturing facility closure, analyze the transfer pricing implications:
    (i)   What is the "exit charge" concept under OECD Guidelines?
    (ii)  What functions, assets, and risks are being transferred?
    (iii) What compensation is appropriate?
    (10 marks)

(b) Evaluate the transfer pricing for the new manufacturing arrangement with Vietnam.
    Consider:
    (i)   The appropriate transfer pricing method
    (ii)  The impact of increased capacity utilization
    (iii) Comparability analysis for the cost-plus markup
    (8 marks)

(c) What are the key transfer pricing documentation requirements for a business restructuring
    of this magnitude? What specific items should be included?
    (7 marks)
"""
    add_paragraph(doc, scenario)

    add_heading(doc, 'QUESTION 2 (25 marks)', level=2)

    question2 = """
Continuing with Artisan Products Restructuring:

(a) Analyze whether the new "principal distributor" model for USA and EU entities is
    appropriate from a transfer pricing perspective. What profit margin should these
    entities achieve post-restructuring?
    (12 marks)

(b) The Italy tax authority has indicated they may challenge the restructuring on grounds
    that an inadequate exit charge was paid to the Italian entity for loss of manufacturing
    functions. Prepare a response defending the group's position or recommending adjustments.
    (8 marks)

(c) Identify other transfer pricing and tax issues associated with this restructuring that
    the group should be aware of, particularly relating to:
    (i)   The IP licensing arrangement
    (ii)  Working capital changes
    (iii) Business restructuring documentation
    (5 marks)
"""
    add_paragraph(doc, question2)
    doc.add_page_break()

def create_exam9_part_a(doc):
    """Mock Exam 9 - Global Services (Business Process Outsourcing)"""
    add_heading(doc, 'QUESTION 1 (25 marks)', level=2)

    scenario = """
SCENARIO - Global Services Provider

Global Services Provider is a multinational company specializing in business process
outsourcing (BPO) for financial services companies. The group provides various back-office
and customer service functions.

GROUP STRUCTURE:

1. Global Services Holdings (Singapore) - Parent company and IP owner
2. Global Services Operations (India) - Primary service delivery
3. Global Services Europe (Ireland) - EU service center
4. Global Services Americas (USA) - Americas customer service center
5. Global Services Quality (UK) - Quality assurance and compliance

COMPANY DETAILS:

Global Services Holdings (Singapore):
- Owns all service delivery methodologies, processes, and proprietary tools
- Owns client relationships framework
- Group coordination
- Employees: 20 executives
- Tax rate: 5% (concessional rate)

Global Services Operations (India):
- Provides majority of service delivery
- 24/7 operations on customer processes
- Handles volume processing, data entry, claims processing
- Employees: 2,000
- Annual revenue (from charging to other entities): INR 2,500 million
- Annual costs: INR 2,200 million (labor intensive)
- Net profit: INR 300 million
- Tax rate: 30%
- Wages: Average INR 300,000 per employee per annum
- Operating margin: 12%

Global Services Europe (Ireland):
- Manages European client relationships
- Conducts client-facing interactions
- Local quality assurance (EU regulation focus)
- Employees: 250
- Annual revenue from clients: EUR 80 million
- Annual costs paid to India operations: EUR 32 million
- Annual local operating costs: EUR 35 million
- Annual net profit: EUR 13 million
- Tax rate: 12.5%
- Margin on client revenue: 16%

Global Services Americas (USA):
- Manages Americas client relationships
- Customer service and escalation handling
- Regulatory compliance for US clients
- Employees: 300
- Annual revenue from clients: USD 100 million
- Annual costs paid to India operations: USD 35 million
- Annual local operating costs: USD 45 million
- Annual net profit: USD 20 million
- Tax rate: 21%
- Margin on client revenue: 20%

Global Services Quality (UK):
- Conducts quality audits and reviews
- Develops compliance protocols
- Tests and validates processes
- Employees: 100
- Annual revenue (internal service charges): GBP 8 million
- Annual costs: GBP 6 million
- Net profit: GBP 2 million
- Tax rate: 25%

SERVICE DELIVERY MODEL:

Client enters contract with Global Services Europe or Americas entities for BPO services.
Revenue flows as:
1. Client pays Global Services Europe/Americas
2. Global Services Europe/Americas pay Global Services Operations (India) for service delivery
3. Global Services Operations manages daily operations
4. Global Services Quality conducts quality assurance
5. Profit retained by client-facing entities (Europe and Americas)

REQUIRED:

(a) Perform a functional analysis for:
    (i)   Global Services Operations (India)
    (ii)  Global Services Europe (Ireland)
    (iii) Global Services Americas (USA)

    Identify functions, assets, and risks for each. Does each entity earn an appropriate
    profit for its functions?
    (10 marks)

(b) Evaluate the transfer pricing arrangement whereby Global Services India is paid on a
    cost-based model while Europe and Americas retain significant margins:
    (i)   What transfer pricing method is being used?
    (ii)  Is it appropriate for this type of service?
    (iii) Are the profit margins consistent with arm's length pricing?
    (10 marks)

(c) What comparables or benchmarking data would you require to support these transfer
    pricing positions? Discuss the challenges in finding relevant comparables for BPO
    services.
    (5 marks)
"""
    add_paragraph(doc, scenario)

    add_heading(doc, 'QUESTION 2 (25 marks)', level=2)

    question2 = """
Continuing with Global Services Provider:

(a) The Indian tax authority has challenged the transfer pricing, arguing that Global
    Services Operations should receive a higher allocation of profits given its role in
    service delivery. Prepare a complete response defending the group's current position or
    recommending changes.
    (12 marks)

(b) Global Services Europe is considering expanding its role by taking over some quality
    assurance functions currently handled by UK. Discuss the transfer pricing implications:
    (i)   Impact on the cost paid to India operations
    (ii)  Impact on margins of each entity
    (iii) Defensibility of the revised structure
    (8 marks)

(c) Recommend a revised transfer pricing structure that you believe would be more robust to
    challenge. Consider using a profit split method and justify your recommendation.
    (5 marks)
"""
    add_paragraph(doc, question2)
    doc.add_page_break()

def create_exam10_part_a(doc):
    """Mock Exam 10 - Green Energy Corp (IP and Licensing)"""
    add_heading(doc, 'QUESTION 1 (25 marks)', level=2)

    scenario = """
SCENARIO - Green Energy Corporation

Green Energy Corporation is a renewable energy company specializing in solar technology
manufacturing and project development. The group has significant IP and technology transfer
arrangements.

GROUP STRUCTURE:

1. Green Energy International (Netherlands) - Parent and technology owner
2. Green Energy Research (Switzerland) - R&D and technology development
3. Green Energy Manufacturing (Mexico) - Panel manufacturing
4. Green Energy Projects (Brazil) - Project development and installation
5. Green Energy Operations (Germany) - EU operations and sales

COMPANY PROFILES:

Green Energy International (Netherlands):
- Holds all patents and proprietary technologies
- Solar panel efficiency technology (proprietary manufacturing process)
- Project development methodologies
- Brand and market positioning
- Employees: 15
- Tax rate: 21%
- IP portfolio value: USD 500 million (estimated)

Green Energy Research (Switzerland):
- Conducts R&D on next-generation solar technology
- Product improvement and cost reduction research
- Employees: 50 engineers and scientists
- Annual R&D costs: CHF 30 million
- Conducts research on cost-reimbursement basis
- No product sales
- Tax rate: 13%
- Research funded by:
  * 20% from operating company sales
  * 80% from technology licensing fees

Green Energy Manufacturing (Mexico):
- Manufactures solar panels using proprietary technology
- Licensed manufacturing rights from parent company
- Produces 500,000 units per annum
- Manufacturing cost per unit: USD 80 (materials, labor, overhead)
- Sells to operating entities at: USD 110 per unit
- Margin: USD 30 per unit (markup of 37.5%)
- Annual production volume value: USD 55 million
- Employee count: 300
- Tax rate: 30%
- EBIT: USD 15 million per annum

Green Energy Projects (Brazil):
- Develops, finances, and manages solar projects
- Engineering and permitting
- Project sales to corporate and government clients
- Installation and commissioning
- Operations and maintenance contracts
- Employees: 200
- Annual project revenue: USD 200 million
- Annual project costs: USD 160 million
- Operating margin: 20%
- EBIT: USD 40 million
- Tax rate: 34%

Green Energy Operations (Germany):
- Sales and distribution hub for Europe
- Sells solar panels to customers in Europe
- Installation services (subcontracted)
- Warranty and service
- Employees: 100
- Annual panel purchases: 100,000 units at USD 110 per unit = USD 11 million
- Annual panel sales revenue: EUR 15 million (approximately USD 16.5 million)
- Annual service revenue: EUR 2 million
- Annual operating costs: EUR 5 million
- Net profit: EUR 3.5 million
- Tax rate: 30%

IP LICENSING ARRANGEMENTS:

1. Manufacturing License (Green Energy Research to Manufacturing):
   - Royalty: 2% of manufacturing revenue
   - Annual royalty: USD 1.1 million

2. Manufacturing License (Green Energy International to Manufacturing):
   - Royalty: 3% of manufacturing revenue
   - Annual royalty: USD 1.65 million

3. Project Development License (Green Energy International to Projects):
   - Royalty: 2% of project revenue
   - Annual royalty: USD 4 million

4. Technology License (Green Energy International to Operations):
   - Royalty: 1% of sales revenue
   - Annual royalty: EUR 0.17 million

ADDITIONAL INFORMATION:
- Solar panel manufacturing is highly competitive globally
- Market royalty rates for solar technology typically range 2-4%
- Green Energy's technology provides 3-5% efficiency advantage
- Next-generation technology under development expected to provide 8-10% advantage
- No technology licensing APAs in place
- Company has not been subject to transfer pricing audit

REQUIRED:

(a) Perform DEMPE analysis for the solar panel technology:
    (i)   Who develops the technology?
    (ii)  Who enhances/improves it?
    (iii) Who maintains it?
    (iv)  Who protects it (IP protection)?
    (v)   Who exploits it (generates revenue)?

    Based on DEMPE, are the royalty rates appropriate and allocated to the right entities?
    (12 marks)

(b) Evaluate the transfer pricing for the manufacturing arrangement in Mexico:
    (i)   Is the cost-plus 37.5% markup appropriate?
    (ii)  What factors support or challenge this markup?
    (iii) Are there benchmarking or comparable data you would require?
    (10 marks)

(c) Identify transfer pricing risks in this structure relating to IP licensing under BEPS
    Actions 8-10.
    (3 marks)
"""
    add_paragraph(doc, scenario)

    add_heading(doc, 'QUESTION 2 (25 marks)', level=2)

    question2 = """
Continuing with Green Energy Corporation:

(a) The Brazilian tax authority has proposed that Green Energy Projects should be classified
    as a "limited risk distributor" earning only routine margins rather than residual
    profits. Analyze and respond to this position:
    (i)   What functions does Green Energy Projects perform?
    (ii)  What risks does it assume?
    (iii) Is the "limited risk distributor" characterization appropriate?
    (iv)  What profit margin should it earn?
    (10 marks)

(b) Green Energy is developing next-generation solar technology (8-10% efficiency advantage).
    Discuss transfer pricing issues that might arise if:
    (i)   The new technology is developed in Switzerland (Research entity)
    (ii)  Current manufacturing in Mexico uses existing technology
    (iii) The new technology will be licensed to Mexico in Year 3
    (12 marks)

(c) Recommend a revised IP licensing and transfer pricing structure that would be robust to
    transfer pricing challenge and consistent with BEPS Actions 8-10.
    (3 marks)
"""
    add_paragraph(doc, question2)
    doc.add_page_break()

def create_part_b(doc, exam_num):
    """Create Part B - Choose One Scenario Question"""
    add_heading(doc, 'PART B: CHOOSE ONE SCENARIO QUESTION (20 marks)', level=1)
    add_paragraph(doc, 'Answer ONE question from the two offered.', italic=True)
    doc.add_paragraph()

    add_heading(doc, 'QUESTION 3 (20 marks)', level=2)

    if exam_num % 2 == 1:
        # Odd exam numbers get restructuring Q
        question3 = """
Your client is a multinational group that is undertaking a significant business restructuring.
The group is considering relocating manufacturing operations from a developed country (Current
Location) to a developing country (New Location) to reduce costs. The current manufacturer
achieves a net profit margin of 15% on manufacturing activities. The new location can
manufacture the same product at 35% lower cost.

(a) Discuss the transfer pricing issues arising from this relocation, specifically:
    (i)   Exit charge implications for the current manufacturer
    (ii)  Transfer pricing methodology for the new manufacturer
    (iii) Working capital and inventory adjustment implications
    (iv)  The impact of the cost reduction on future transfer prices
    (8 marks)

(b) Recommend a defensible approach to pricing this restructuring that balances commercial
    rationale with transfer pricing compliance. Consider any APAs or advance notice to tax
    authorities that might be appropriate.
    (12 marks)
"""
    else:
        # Even exam numbers get PE Q
        question3 = """
Your client operates a substantial business in a foreign country where it has stationed
personnel to manage operations and support local sales. The foreign tax authority has indicated
it believes your client has established a Permanent Establishment (PE) under Article 5 of the
relevant tax treaty and is pursuing a transfer pricing audit on the basis that profits should
be attributed to the PE.

(a) Analyze the PE issue comprehensively:
    (i)   What is the applicable test under Article 5?
    (ii)  What factors would support the existence of a PE?
    (iii) What factors would militate against PE status?
    (iv)  How would you structure the operations to minimize PE risk?
    (8 marks)

(b) If a PE is determined to exist, analyze the transfer pricing implications:
    (i)   How should profits be attributed using the Authorized OECD Approach?
    (ii)  What functions and assets would you attribute to the PE?
    (iii) What profit margin would be appropriate?
    (12 marks)
"""

    add_paragraph(doc, question3)

    doc.add_paragraph()
    add_heading(doc, 'QUESTION 4 (20 marks)', level=2)

    if exam_num % 2 == 1:
        # Odd exam numbers get PE Q
        question4 = """
Your client operates an international service delivery model where a service provider based
in a low-tax jurisdiction provides technical services to operating entities in higher-tax
jurisdictions. The service provider is charging on a cost-plus basis, which results in the
service provider earning a minimal margin while the operating entities earn substantial
profits. A tax authority in one of the higher-tax jurisdictions has challenged whether the
service provider should be earning a higher margin and whether the service provider might
constitute a Permanent Establishment.

(a) Analyze the transfer pricing issue:
    (i)   What is the current transfer pricing methodology?
    (ii)  Is the cost-plus basis appropriate for these services?
    (iii) What alternative transfer pricing methodologies might be considered?
    (iv)  What would be an arm's length profit margin for the service provider?
    (8 marks)

(b) Evaluate the PE risk and recommend a structure that minimizes PE exposure while
    maintaining transfer pricing defensibility:
    (i)   What functions and assets of the service provider increase PE risk?
    (ii)  How could the structure be modified to reduce PE exposure?
    (iii) What documentation should be prepared?
    (12 marks)
"""
    else:
        # Even exam numbers get restructuring Q
        question4 = """
Your client is a multinational group that is considering a business restructuring involving
the transfer of intellectual property (IP) from one jurisdiction to another. The current owner
of the IP is in a high-tax jurisdiction, and the client is considering transferring the IP to
a lower-tax jurisdiction to generate IP licensing revenues there. The IP transfer will occur
at a time when the technology is being enhanced and new patents are being developed.

(a) Discuss the transfer pricing issues arising from this IP transfer:
    (i)   How should the transfer price be determined?
    (ii)  What BEPS Action 8-10 principles apply to this situation?
    (iii) What role does the DEMPE framework play in analyzing this transfer?
    (iv)  How should the allocation of ongoing R&D costs be handled?
    (8 marks)

(b) Recommend a transfer pricing position for the IP transfer and the subsequent licensing
    arrangements that would be robust to challenge, considering both the transfer price and
    the ongoing royalty rates:
    (i)   What valuation method would you use?
    (ii)  What would be an appropriate royalty rate for subsequent licenses?
    (iii) What documentation is essential?
    (12 marks)
"""

    add_paragraph(doc, question4)
    doc.add_page_break()

def create_part_c(doc):
    """Create Part C - Choose Two Short Questions"""
    add_heading(doc, 'PART C: CHOOSE TWO SHORT QUESTIONS (15 marks each)', level=1)
    add_paragraph(doc, 'Answer TWO questions from the five offered. Each question is worth 15 marks.', italic=True)
    doc.add_paragraph()

    questions_c = [
        ("QUESTION 5", """
INTRAGROUP SERVICES AND SERVICE HUBS

Discuss the key transfer pricing issues for intragroup services, including:
(a) How to distinguish between:
    (i)   Genuine intragroup services subject to transfer pricing
    (ii)  Shareholder activities (not subject to charge)
    (iii) Duplicative services (not subject to charge)
    (8 marks)

(b) The transfer pricing methodology appropriate for:
    (i)   Administrative and back-office services
    (ii)  Technical and professional services
    (iii) Management and coordination services
    (7 marks)
"""),
        ("QUESTION 6", """
INTELLECTUAL PROPERTY AND DEMPE FRAMEWORK

Explain the DEMPE framework (Development, Enhancement, Maintenance, Protection, Exploitation)
and its application to transfer pricing, addressing:
(a) How BEPS Actions 8-10 have changed transfer pricing analysis for intellectual property
    (8 marks)

(b) How DEMPE analysis should influence:
    (i)   The allocation of R&D costs among group entities
    (ii)  The transfer pricing methodology for IP licensing
    (iii) The profit allocation between IP owners and users
    (7 marks)
"""),
        ("QUESTION 7", """
ADVANCE PRICING AGREEMENTS

Discuss Advance Pricing Agreements (APAs), including:
(a) When it is appropriate for a multinational group to seek an APA:
    (i)   Advantages of APAs
    (ii)  Disadvantages of APAs
    (iii) Bilateral vs. unilateral APAs
    (8 marks)

(b) The key issues that should be addressed in an APA covering:
    (i)   Intercompany transactions to be covered
    (ii)  Transfer pricing methodology and supporting documentation
    (iii) The term of the APA and adjustment mechanisms
    (7 marks)
"""),
        ("QUESTION 8", """
TRANSFER PRICING DOCUMENTATION AND COMPLIANCE

Discuss transfer pricing documentation requirements under BEPS Action 13, including:
(a) The three-tiered approach to transfer pricing documentation:
    (i)   Master File requirements
    (ii)  Local File requirements
    (iii) Country-by-Country Reporting requirements
    (8 marks)

(b) Best practices for preparation of transfer pricing documentation, including:
    (i)   Benchmarking and comparables analysis
    (ii)  Functional analysis and FAR analysis
    (iii) Supporting contemporaneous documentation
    (7 marks)
"""),
        ("QUESTION 9", """
TRANSFER PRICING AND TAX TREATIES

Discuss the interaction between transfer pricing and tax treaties, addressing:
(a) Key articles in the OECD Model Tax Convention relevant to transfer pricing:
    (i)   Associated Enterprises (Article 9)
    (ii)  Non-Discrimination (Article 24)
    (iii) Mutual Agreement Procedure (Article 25)
    (8 marks)

(b) How a tax treaty can:
    (i)   Provide protection against double taxation
    (ii)  Facilitate Mutual Agreement Procedures
    (iii) Influence transfer pricing methodology selection
    (7 marks)
"""),
    ]

    for question_num, question_text in enumerate(questions_c, 5):
        if question_num > 6:
            add_heading(doc, questions_c[question_num - 5][0], level=2)
            add_paragraph(doc, questions_c[question_num - 5][1])

    # Just add all 5 questions for Part C
    for i, (question_title, question_text) in enumerate(questions_c):
        add_heading(doc, question_title, level=2)
        add_paragraph(doc, question_text)
        if i < len(questions_c) - 1:
            doc.add_paragraph()

def create_complete_exam(exam_num):
    """Create a complete mock exam"""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Create exam content
    create_exam_header(doc, exam_num)
    create_part_a(doc, exam_num)
    create_part_b(doc, exam_num)
    create_part_c(doc)

    # Save document
    output_path = f'/home/user/ADITER/MockExam/MockExam_{exam_num:02d}.docx'
    doc.save(output_path)
    return output_path

# Main execution
if __name__ == '__main__':
    output_files = []
    for exam_num in range(1, 11):
        try:
            output_path = create_complete_exam(exam_num)
            output_files.append(output_path)
            print(f'Created: {output_path}')
        except Exception as e:
            print(f'Error creating exam {exam_num}: {e}')

    print(f'\n✓ Successfully created {len(output_files)} mock exams')
    print(f'Saved to: /home/user/ADITER/MockExam/')
