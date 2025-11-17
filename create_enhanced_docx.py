#!/usr/bin/env python3
"""
Script to create enhanced Word documents from markdown content
Following QA Enhancement Protocol formatting standards
Enhanced with robust markdown parsing and formatting preservation
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def create_enhanced_document(markdown_content, title, output_filename):
    """
    Create a professionally formatted Word document from markdown content

    Args:
        markdown_content: The enhanced markdown content
        title: Document title
        output_filename: Output .docx filename
    """
    doc = Document()

    # Set up styles
    setup_styles(doc)

    # Parse and add content
    parse_markdown_to_docx(doc, markdown_content, title)

    # Save document
    doc.save(output_filename)
    print(f"✅ Created: {output_filename}")

def setup_styles(doc):
    """Configure document styles per QA Enhancement Protocol"""
    styles = doc.styles

    # Configure Heading 1 (Document title)
    h1 = styles['Heading 1']
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.name = 'Calibri'

    # Configure Heading 2 (Main sections)
    h2 = styles['Heading 2']
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.name = 'Calibri'

    # Configure Heading 3 (Subsections)
    h3 = styles['Heading 3']
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.name = 'Calibri'

    # Configure Heading 4 (Sub-subsections)
    h4 = styles['Heading 4']
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.name = 'Calibri'

    # Configure Normal style
    normal = styles['Normal']
    normal.font.size = Pt(11)
    normal.font.name = 'Calibri'
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

def parse_markdown_to_docx(doc, content, title):
    """Parse markdown and add to document with proper formatting"""

    try:
        # Add title
        doc.add_heading(title, level=1)

        lines = content.split('\n')
        i = 0

        while i < len(lines):
            try:
                line = lines[i].rstrip()

                # Skip empty lines at start
                if not line and i == 0:
                    i += 1
                    continue

                # Skip title line if it matches our title
                if line.startswith('# ') and title in line:
                    i += 1
                    continue

                # Horizontal rules (---)
                if line.strip() == '---' or line.strip() == '***' or line.strip() == '___':
                    add_horizontal_rule(doc)
                    i += 1
                    continue

                # Headers
                if line.startswith('#### '):
                    doc.add_heading(line[5:], level=4)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)

                # Tables
                elif line.startswith('|') and i + 1 < len(lines) and lines[i+1].startswith('|'):
                    i = add_table(doc, lines, i)
                    continue

                # Lists (including nested)
                elif line.startswith('- ') or line.startswith('  -') or line.startswith('    -') or re.match(r'^(\s*)\d+\.\s', line):
                    i = add_list(doc, lines, i)
                    continue

                # Code blocks
                elif line.startswith('```'):
                    i = add_code_block(doc, lines, i)
                    continue

                # Regular paragraphs
                elif line:
                    p = doc.add_paragraph()
                    add_formatted_text(p, line)

                i += 1

            except Exception as e:
                # Error recovery: skip problematic line and continue
                print(f"⚠️  Warning: Error processing line {i}: {str(e)}")
                i += 1
                continue

    except Exception as e:
        print(f"❌ Error during document parsing: {str(e)}")
        raise

def add_horizontal_rule(doc):
    """Add a horizontal rule (line) to the document"""
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_before = Pt(6)
    p_format.space_after = Pt(6)

    # Add bottom border to create horizontal line
    p_element = p._element
    p_pr = p_element.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')

    p_bdr.append(bottom)
    p_pr.append(p_bdr)

def add_table(doc, lines, start_idx):
    """Parse and add a markdown table with cell formatting support"""
    try:
        table_lines = []
        i = start_idx

        # Collect table lines
        while i < len(lines) and lines[i].startswith('|'):
            table_lines.append(lines[i])
            i += 1

        if len(table_lines) < 2:
            return i

        # Parse table
        headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
        rows = []

        for line in table_lines[2:]:  # Skip header separator
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)

        # Create table
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Light Grid Accent 1'

        # Add headers with formatting
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = ''  # Clear default text
            # Parse markdown in header
            add_formatted_text(cell.paragraphs[0], header)
            # Make headers bold
            for run in cell.paragraphs[0].runs:
                run.font.bold = True

        # Add rows with formatting
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = ''  # Clear default text
                # Parse markdown formatting in cell
                add_formatted_text(cell.paragraphs[0], cell_data)

        doc.add_paragraph()  # Space after table
        return i

    except Exception as e:
        print(f"⚠️  Warning: Error parsing table: {str(e)}")
        return start_idx + 1

def add_list(doc, lines, start_idx):
    """Parse and add a list with multi-level nesting support"""
    try:
        i = start_idx

        while i < len(lines):
            line = lines[i].rstrip()

            if not line:
                break

            # Calculate indent level based on leading spaces
            indent_level = 0
            stripped = line.lstrip()
            leading_spaces = len(line) - len(stripped)

            # 2 spaces = 1 indent level
            if leading_spaces >= 2:
                indent_level = leading_spaces // 2

            # Bullet list
            if stripped.startswith('- '):
                text = stripped[2:]
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_text(p, text)
                # Set indent level
                if indent_level > 0:
                    p.paragraph_format.left_indent = Inches(0.25 * indent_level)

            # Numbered list
            elif re.match(r'^\d+\.\s', stripped):
                match = re.match(r'^\d+\.\s(.+)', stripped)
                if match:
                    text = match.group(1)
                    p = doc.add_paragraph(style='List Number')
                    add_formatted_text(p, text)
                    # Set indent level
                    if indent_level > 0:
                        p.paragraph_format.left_indent = Inches(0.25 * indent_level)
            else:
                break

            i += 1

        return i

    except Exception as e:
        print(f"⚠️  Warning: Error parsing list: {str(e)}")
        return start_idx + 1

def add_code_block(doc, lines, start_idx):
    """Add a code block with monospace font and shading"""
    try:
        i = start_idx + 1
        code_lines = []

        # Collect code block lines
        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i])
            i += 1

        # Detect if this is an ASCII diagram (contains arrows or diagram symbols)
        is_ascii_diagram = any(
            arrow in line for line in code_lines
            for arrow in ['→', '←', '↓', '↑', '├', '└', '│', '─', '↔']
        )

        # Add as formatted code block with shading
        for code_line in code_lines:
            p = doc.add_paragraph()

            # Apply monospace font for code/calculations/diagrams
            run = p.add_run(code_line)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

            # Add shading to paragraph
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F0F0F0')  # Light gray background
            p._element.get_or_add_pPr().append(shading_elm)

            # Tighter spacing for code blocks
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

        # Add spacing after code block
        doc.add_paragraph()

        return i + 1 if i < len(lines) else i

    except Exception as e:
        print(f"⚠️  Warning: Error parsing code block: {str(e)}")
        return start_idx + 1

def add_formatted_text(paragraph, text):
    """
    Add text with robust markdown formatting support:
    - **bold**, *italic*, `code`, ***bold+italic***
    - Nested formatting
    - Special symbols preserved
    """
    try:
        # Parse inline formatting with proper tokenization
        tokens = tokenize_markdown(text)

        for token in tokens:
            token_type = token['type']
            content = token['content']

            if token_type == 'bold':
                run = paragraph.add_run(content)
                run.font.bold = True
            elif token_type == 'italic':
                run = paragraph.add_run(content)
                run.font.italic = True
            elif token_type == 'bold_italic':
                run = paragraph.add_run(content)
                run.font.bold = True
                run.font.italic = True
            elif token_type == 'code':
                run = paragraph.add_run(content)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                # Light background for inline code
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F5F5F5')
                run._element.get_or_add_rPr().append(shading_elm)
            else:  # plain text
                paragraph.add_run(content)

    except Exception as e:
        # Fallback: add as plain text if parsing fails
        print(f"⚠️  Warning: Error parsing formatted text, using plain text: {str(e)}")
        paragraph.add_run(text)

def tokenize_markdown(text):
    """
    Tokenize markdown text into formatting elements
    Returns list of tokens: [{'type': 'bold', 'content': 'text'}, ...]
    """
    tokens = []
    i = 0

    while i < len(text):
        # Check for bold+italic (***text***)
        if text[i:i+3] == '***':
            end = text.find('***', i + 3)
            if end != -1:
                tokens.append({'type': 'bold_italic', 'content': text[i+3:end]})
                i = end + 3
                continue

        # Check for bold (**text**)
        if text[i:i+2] == '**':
            end = text.find('**', i + 2)
            if end != -1:
                tokens.append({'type': 'bold', 'content': text[i+2:end]})
                i = end + 2
                continue

        # Check for italic (*text*)
        if text[i:i+1] == '*' and (i == 0 or text[i-1] != '*'):
            end = text.find('*', i + 1)
            # Make sure it's not part of **
            if end != -1 and (end + 1 >= len(text) or text[end+1] != '*'):
                tokens.append({'type': 'italic', 'content': text[i+1:end]})
                i = end + 1
                continue

        # Check for inline code (`code`)
        if text[i:i+1] == '`':
            end = text.find('`', i + 1)
            if end != -1:
                tokens.append({'type': 'code', 'content': text[i+1:end]})
                i = end + 1
                continue

        # Plain text - find next special character
        next_special = len(text)
        for special in ['*', '`']:
            pos = text.find(special, i)
            if pos != -1 and pos < next_special:
                next_special = pos

        if next_special == len(text):
            tokens.append({'type': 'text', 'content': text[i:]})
            break
        else:
            if next_special > i:
                tokens.append({'type': 'text', 'content': text[i:next_special]})
            i = next_special

    return tokens


if __name__ == '__main__':
    print("Enhanced Word Document Creator")
    print("=" * 50)
