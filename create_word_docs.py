#!/usr/bin/env python3
"""
Create professional Word documents from markdown files
Following QA Enhancement Protocol formatting standards
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from datetime import datetime

def add_shading(cell, color):
    """Add shading to table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_word_document(md_file, title, output_file):
    """
    Create a professionally formatted Word document from markdown

    Args:
        md_file: Path to markdown file
        title: Document title
        output_file: Output .docx filename
    """
    print(f"Processing: {md_file}")

    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set up document styles
    setup_document_styles(doc)

    # Add title
    title_para = doc.add_heading(title, level=1)

    # Parse and add content
    lines = content.split('\n')
    i = 0
    in_table = False
    table_lines = []
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip the main title if it matches
        if line.startswith('# ') and title.replace(' - ', ' ').lower() in line.lower():
            i += 1
            continue

        # Handle code blocks
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # End of code block - add it
                add_code_block(doc, code_lines)
                in_code_block = False
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Handle tables
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # End of table
            add_markdown_table(doc, table_lines)
            in_table = False
            table_lines = []
            continue

        # Handle headers
        if line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)

        # Handle bullet lists
        elif line.startswith('- '):
            add_formatted_paragraph(doc, line[2:], style='List Bullet')

        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line):
            match = re.match(r'^\d+\.\s(.+)', line)
            if match:
                add_formatted_paragraph(doc, match.group(1), style='List Number')

        # Handle horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph()  # Just add spacing

        # Handle regular paragraphs
        elif line.strip():
            add_formatted_paragraph(doc, line)

        # Handle empty lines
        else:
            # Don't add too many blank paragraphs
            pass

        i += 1

    # Handle any remaining table
    if in_table and table_lines:
        add_markdown_table(doc, table_lines)

    # Save document
    doc.save(output_file)
    print(f"✅ Created: {output_file}")
    return output_file

def setup_document_styles(doc):
    """Configure document styles per QA Enhancement Protocol"""
    styles = doc.styles

    # Heading 1 - Document title (16pt, bold)
    h1 = styles['Heading 1']
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.name = 'Calibri'
    h1.font.color.rgb = RGBColor(0, 0, 0)

    # Heading 2 - Main sections (14pt, bold)
    h2 = styles['Heading 2']
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.name = 'Calibri'
    h2.font.color.rgb = RGBColor(0, 0, 0)

    # Heading 3 - Subsections (12pt, bold)
    h3 = styles['Heading 3']
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.name = 'Calibri'
    h3.font.color.rgb = RGBColor(0, 0, 0)

    # Heading 4 - Sub-subsections (11pt, bold)
    h4 = styles['Heading 4']
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.name = 'Calibri'
    h4.font.color.rgb = RGBColor(0, 0, 0)

    # Normal - Body text (11pt, Calibri, 1.15 spacing, 6pt after)
    normal = styles['Normal']
    normal.font.size = Pt(11)
    normal.font.name = 'Calibri'
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_formatted_paragraph(doc, text, style='Normal'):
    """Add paragraph with bold/italic formatting"""
    p = doc.add_paragraph(style=style)

    # Parse bold (**text**) and italic (*text*)
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = p.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            # Italic text
            run = p.add_run(part[1:-1])
            run.font.italic = True
        elif part:
            p.add_run(part)

def add_markdown_table(doc, table_lines):
    """Parse markdown table and add to document"""
    if len(table_lines) < 2:
        return

    # Parse header
    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]

    # Skip separator line (second line with |---|---|)
    data_lines = [line for line in table_lines[2:] if line.strip() and '|' in line]

    if not headers or not data_lines:
        return

    # Create table
    num_cols = len(headers)
    num_rows = len(data_lines)

    table = doc.add_table(rows=num_rows + 1, cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Add headers
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        # Bold header text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        # Shade header
        add_shading(cell, 'D9E2F3')

    # Add data rows
    for row_idx, line in enumerate(data_lines):
        cells_data = [cell.strip() for cell in line.split('|')[1:-1]]
        for col_idx, cell_data in enumerate(cells_data):
            if col_idx < num_cols:
                table.rows[row_idx + 1].cells[col_idx].text = cell_data

    # Add spacing after table
    doc.add_paragraph()

def add_code_block(doc, code_lines):
    """Add code block with shading"""
    if not code_lines:
        return

    for line in code_lines:
        p = doc.add_paragraph(line)
        # Make it look like code
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        # Add light gray background (simulated with spacing)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)

def main():
    """Main function to create all Word documents"""
    today = datetime.now().strftime('%Y-%m-%d')

    files_to_process = [
        {
            'input': '/home/user/ADITER/generated_notes/Fundamental_tax_issues_International_investment.md',
            'corrected': None,  # Will use original with QA knowledge
            'title': 'International Investment - Tax Considerations in Energy Resources',
            'output': f'/home/user/ADITER/generated_notes/Fundamental_tax_issues_International_investment_Enhanced_{today}.docx'
        },
        {
            'input': '/home/user/ADITER/generated_notes/Fundamental_tax_issues_International_income_flows.md',
            'corrected': '/home/user/ADITER/generated_notes/Fundamental_tax_issues_International_income_flows_CORRECTED.md',
            'title': 'International Income Flows - Tax Considerations in Energy Resources',
            'output': f'/home/user/ADITER/generated_notes/Fundamental_tax_issues_International_income_flows_Enhanced_{today}.docx'
        },
        {
            'input': '/home/user/ADITER/generated_notes/Fundamental_tax_issues_Tax_treaties.md',
            'corrected': '/home/user/ADITER/generated_notes/Fundamental_tax_issues_Tax_treaties_CORRECTED.md',
            'title': 'Tax Treaties in Energy Resources',
            'output': f'/home/user/ADITER/generated_notes/Fundamental_tax_issues_Tax_treaties_Enhanced_{today}.docx'
        },
        {
            'input': '/home/user/ADITER/generated_notes/Fundamental_tax_issues_Transfer_pricing.md',
            'corrected': '/home/user/ADITER/generated_notes/Fundamental_tax_issues_Transfer_pricing_CORRECTED.md',
            'title': 'Transfer Pricing in Energy Resources',
            'output': f'/home/user/ADITER/generated_notes/Fundamental_tax_issues_Transfer_pricing_Enhanced_{today}.docx'
        }
    ]

    print("=" * 70)
    print("QA ENHANCED WORD DOCUMENT GENERATOR")
    print("=" * 70)
    print()

    created_files = []

    for file_info in files_to_process:
        # Use corrected file if it exists, otherwise use original
        import os
        if file_info['corrected'] and os.path.exists(file_info['corrected']):
            input_file = file_info['corrected']
            print(f"Using QA-corrected version: {os.path.basename(input_file)}")
        else:
            input_file = file_info['input']
            print(f"Using original: {os.path.basename(input_file)}")

        try:
            output = create_word_document(
                input_file,
                file_info['title'],
                file_info['output']
            )
            created_files.append(output)
            print()
        except Exception as e:
            print(f"❌ Error creating {file_info['output']}: {e}")
            print()

    print("=" * 70)
    print(f"SUMMARY: Created {len(created_files)} Word documents")
    print("=" * 70)
    for f in created_files:
        print(f"  📄 {f}")
    print()

if __name__ == '__main__':
    main()
