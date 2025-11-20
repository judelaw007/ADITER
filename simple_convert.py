#!/usr/bin/env python3
"""
Simple markdown to DOCX conversion using python-docx
"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def simple_convert_md_to_docx(md_path, docx_path):
    """Simple conversion - just add markdown content with basic formatting"""
    doc = Document()

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Extract title from filename
    filename = os.path.basename(md_path)
    title = filename.replace('.md', '').split('_', 1)[1].replace('_', ' ') if '_' in filename else filename

    # Add title as heading
    doc.add_heading(title, level=1)

    # Add content line by line with simple formatting
    for line in lines:
        line = line.rstrip()

        if not line:
            doc.add_paragraph()
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('  - '):
            p = doc.add_paragraph(line[4:], style='List Bullet 2')
        elif line.startswith('    - '):
            p = doc.add_paragraph(line[6:], style='List Bullet 3')
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)
    return True

def main():
    input_dir = "/home/user/ADITER/generated_notes/12_Transfer_pricing/quality_checked"
    md_files = sorted([f for f in os.listdir(input_dir) if f.endswith('.md')])

    print(f"Found {len(md_files)} markdown files\n")

    converted = 0
    for idx, md_file in enumerate(md_files, 1):
        input_path = os.path.join(input_dir, md_file)
        output_filename = md_file.replace('.md', '.docx')
        output_path = os.path.join(input_dir, output_filename)

        # Skip if already exists
        if os.path.exists(output_path):
            print(f"[{idx}/{len(md_files)}] {md_file} - Already exists")
            continue

        print(f"[{idx}/{len(md_files)}] Converting: {md_file}...", end='', flush=True)

        try:
            simple_convert_md_to_docx(input_path, output_path)
            print(" ✅ Done")
            converted += 1
        except Exception as e:
            print(f" ❌ Error: {e}")

    print(f"\nConverted {converted} files")

if __name__ == '__main__':
    main()
